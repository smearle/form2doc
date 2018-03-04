import sys
from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
import re
import os

def get_forms(form_path,inputs_to_form_entries={}):
    form = PdfFileReader(form_path).getFields()
    worker_inputs_to_form_entries = {}
    for k in inputs_to_form_entries.keys():
        try:
            # print(form[k])
            # print(inputs_to_form_entries[k])
            worker_inputs_to_form_entries[k]=form[k]
        except:
            print('Form '+ form_path +' does not contain '+k)
    return worker_inputs_to_form_entries

# Loop through forms supplied as arguments
def form2doc(filedir):
    if os.path.isdir(filedir):
        # print('issa dir')
        for j in os.listdir(filedir):
            form2doc(j)
    elif os.path.isfile(filedir) and filedir.endswith('.pdf'):
        # extract information from PDF form
        form = PdfFileReader(filedir).getFields()
        # keys = ['JOB NAMECOMPANY','FULL NAME AS APPEARS ON PASSPORT','SEX','DATE OF BIRTH MMDDYYYY','SOCIAL SECURITY NUMBER','ADDRESS','TELEPHONE NUMBER','POSITION  JOB TITLE','EXPECTED TOTAL WAGES FOR THIS JOB','COUNTRY OF BIRTH']
        # info = [form[k] for k in keys]
        fullname = form['FULL NAME AS APPEARS ON PASSPORT']['/V']
        position = form['POSITION  JOB TITLE']['/V']
        jobname_company = form['JOB NAMECOMPANY']['/V']
        jobname_company = re.split('(/)', jobname_company)
        # first element is the job name, second is the company
        jobname=jobname_company[0]
        company= jobname_company[-1]
        sex = form['SEX']['/V']
        dob = form['DATE OF BIRTH MMDDYYYY']['/V']
        sin = form['SOCIAL SECURITY NUMBER']['/V']
        addy= form['ADDRESS']['/V']
        tel = form['TELEPHONE NUMBER']['/V']
        shootdates = form['DATE OF JOB SHOOTING']['/V']
        shootdates = re.split('(to|-)',shootdates)
        startdate = shootdates[0].strip()
        enddate = shootdates[-1].strip()
        birthdate = form['DATE OF BIRTH MMDDYYYY']['/V']
        immigration_processing = form['IMMIGRATION PROCESSING AT']['/V']
        birthcountry = form['COUNTRY OF BIRTH']['/V']
        sex = form['SEX']['/V']
        sex_entry_to_gender={'male':0,'m':0, 'man':0, 'homme':0,'masc':0,
        'female':1,'f':1,'woman':1,'femme':1}
        possessive = ['His','Her','Their']
        subject = ['He','She','They']
        pers_object =['Him','Her','They']
        sex = sex.lower().replace('-','').replace(' ','').encode('ascii')
        if sex in sex_entry_to_gender.keys():
            d = sex_entry_to_gender[sex]
        else:
            d = 2
        possessive = possessive[d]
        subject = subject[d]
        pers_object = pers_object[d]


        # The replacement dictionary
        replacement = {'[name, title]': fullname +', '+position,
        '[name]': fullname,
        '[name here]': fullname,
        '[client]': fullname,
        '[visitor name]' : fullname,
        '[visitor name' : fullname,

        'dob m/d/y': birthdate,

        'her':possessive,
        'her/':possessive,
        '[her':possessive,
        '[her/':possessive,
        '[her/his]':possessive,

        '[she/he]':subject,

        '[position]': position,
        '[foreign production company]':'[FOREIGN PRODUCTION COMPANY]',
        '[foreign production ': '[FOREIGN PRODUCTION COMPANY]',

        '[producer]': company,
        '[city]':immigration_processing,
        '[name of job]': jobname,
        '[production]': jobname,
        '[job name]': jobname,
        '[job]': jobname,
        '[start date]':startdate,
        '[end date]': enddate,
        '[budget]': '[BUDGET]',
        '[nationality]': birthcountry,
        '[position]': position,
        '[days in canada]': '[DayCalculation]',

        '[agency or client]': '[AGENCY]',
        '[agency]':'[AGENCY]',

        '[shoot dates]':startdate+' and '+enddate,


        u'[description of client\u2019s business]':'BUSINESS DESCRIPTION?',

        '[location]':'Location?',
        '[production dates]': startdate+' and '+enddate+',',
        u'[r186(a)]':u'[r186(a)]'
        }

        # write UNPAID .docx file with client name
        invite = Document('Templates/unpaid.docx')
        for p in invite.paragraphs:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            par_text =''
            text_style=None
            for i in range(len(inline)):
                # print(inline[i].text)
                text_style = inline[i].style
                par_text = par_text + inline[i].text
            boxes = re.findall('(\[.*?\])', par_text)
            match = False
            for box in boxes:
                print(box)
                match=True
                box_0 = box.lower()
                entry = replacement[box_0].strip()
                print(entry + '!')
                par_text = par_text.replace(box, entry).replace('  ',' ')
            # print(par_text)
            if match:
                p.clear()
                p.add_run(par_text, text_style)
        invite.save('Output/'+fullname +' unpaid.docx')

        # Write PAID .docx file w/ client info
        invite2 = Document('Templates/paid.docx')
        for p in invite2.paragraphs:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            par_text=''
            text_style=None
            for i in range(len(inline)):
                text_style = inline[i].style
                par_text = par_text + inline[i].text
            boxes = re.findall('(\[.*?\]|DOB M\/D\/Y)', par_text)
            match = False
            for box in boxes:
                match = True
                print(box)
                box_0 = box.lower()
                entry = replacement[box_0].strip()
                # print(entry)
                if par_text.find(box) >1:
                    if par_text[par_text.find(box)-2] == '.':
                        entry=entry.title()
                par_text = par_text.replace(box, entry).replace('  ',' ')
            # print(par_text + ' END PAR ')
            # print(p.style)
            if match:
                p.clear()
                p.add_run(par_text, text_style)
            # print(p.style)
        invite2.save('Output/'+fullname +' paid.docx')
    else:
        print(filedir+' is either not a valid path, or not a pdf file.')
