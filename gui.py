from appJar import gui
import appJar
from PyPDF2 import PdfFileWriter, PdfFileReader
from docx import Document
import docx
import os
import re
import inspect
import ruleFunctions
import traceback

from form2doc import form2doc
from ruleFunctions import RuleParseError
from tkinter import TclError
from tkinter import Event
from pprint import pprint

class MissingFieldError(Exception):
    pass

out_temp_paths = []

# A dictionary of filepaths to liveDoc objects.
outdocs = {}

instoouts = {}
outstoins = {}

src_dir = os.path.dirname(os.path.realpath(__file__))
# OPC entry letter rules
rulesheet_dirpath = src_dir + '/Rules'
input_templates_dirpath = src_dir + '/Input Templates'
output_templates_dirpath = src_dir + '/Output Templates'

input_dirpath = src_dir + '/Input'
output_dirpath = src_dir + '/Output'

# Dictionary mapping input form paths to PyPDF form-field dictionaries
formpaths_to_formdicts = {}
# Doctionary mapping input form paths to (title to entry) form-field dictionaries
formpaths_to_fielddicts ={}

template_form_path = None
# PdfFileReader object
template_form = None
# Forms dictionary
template_fields = None
# Dictionary of out_templates to lists of rules, each with at least one string pointers
# to potentit\al worker form field, and one to an output template field
# out_template --> [fillable_field(s)] --> [replacement(s)] --> [condition(s)] (optional)
#                                      --> alt [replacement(s)] --> alt [conditions]
outfields2rules = {}

box_ptrn = re.compile('\[.*?\]')

def compileRules():
    rules = [rule.decode('utf8') for rule in app.getAllListItems(app.getTabbedFrameSelectedTab('rulesheets')+'_rules')]
    for rule in rules:
        if rule.replace(' ','') == '':
            rules.delete(rule)

    input_path = app.getTabbedFrameSelectedTab('inputs')

    rule_parser = ruleFunctions.RuleParser(formpaths_to_fielddicts[input_path])

    compiled_rules = [Rule(rule_parser, rulestr) for rulestr in rules]
    return compiled_rules


class Rule(object):
    def __init__(self, rule_parser, rule):
        input_path = app.getTabbedFrameSelectedTab('inputs')
        r_pos = []
        c_pos = []
        op = 0
        cl = 0
        i = 0
        LAST_R = True
        # Find replacees
        replacee_str, replacement_str = '',''
        while rule[i:] != '':
            if rule[i] == '[':
                op+= 1
            elif rule[i] == ']':
                cl += 1
            if op == cl:
                if rule[i:].startswith('replaced by '):
                    replacee_str = rule[:i]
                    replacement_str = rule[i+12:]
                    break
            i+=1
        self.replacees = re.findall(r'\[.*?\]', replacee_str)
        self.replace_codes = re.findall(r'([\w\d]*?)(\{.*?\})',replacement_str)
        self.replacement = replacement_str
        infields = [replacement[1].strip('{}') for replacement in  self.replace_codes]
        funcs = [replacement[0] for replacement in  self.replace_codes]
        for i in range(len( self.replace_codes)):
            func = funcs[i]
            infield = infields[i]
            if func !=None:
                is_valid = False

                for ruleFunc in inspect.getmembers(rule_parser, inspect.ismethod):
                    if func == ruleFunc[0]:
                        is_valid = True
                        argument = infield
                        try:
                            argument = app.getEntry(argument+'_'+input_path)
                            subreplace = getattr(rule_parser, func)(argument)
                        except RuleParseError:
                            print(RuleParseError)
                            subreplace = 'Error'
                        break
                if not is_valid:
                    try:
                        subreplace = app.getEntry(infield+'_'+input_path)
                    except:
                        subreplace = 'Error'
            else:
                subreplace = app.getEntry(infield+'_'+input_path)
            self.replacement = self.replacement.replace(func + '{'+infield+'}', subreplace)
        print('replacements: %s' %  self.replacement)


def saveFormedit():
    f = app.getTabbedFrameSelectedTab('inputs')
    worker_entries = formpaths_to_formdicts[f]
    fields = formpaths_to_fielddicts[f]
    # fields0=fields
    for k in worker_entries.keys():
        fields[k]= app.getEntry(k+f)
        # fields0[k]=''
    outform = PdfFileWriter()
    original_form = PdfFileReader(open(f,'rb'))
    outStream = open('testout.pdf', 'wb')
    # outform.write(outStream)
    # OUTTEST = open('supertesty.pdf','wb')
    # outform.write(OUTTEST)
    # OUTTEST.close()

    for p in range(original_form.getNumPages()):
        outform.addPage(original_form.getPage(p))
    for p in range(outform.getNumPages()):
        # outform.updatePageFormFieldValues(outform.getPage(p), fields0)
        outform.updatePageFormFieldValues(outform.getPage(p), fields)
    outform.write(outStream)
    outform.write(outStream)
    # outform.write(outStream)
    # outform.write(outStream)

    outStream.close()

def saveOutput():
    outpath = app.getTabbedFrameSelectedTab('output_preview')
    outdocs[outpath].save(outpath)


def getFiles(filedir, extension):
    i=0
    for f in filedir:
        if os.path.isfile(f) and (f.endswith('.' + extension)):
            pass
        elif os.path.isdir(f):
            out = []
            for j in os.listdir(f):
                out = out + getFiles([f+'/'+j],extension)
            filedir= filedir[:i]+out+filedir[i+1:]
        else: filedir=filedir[:i]+filedir[i+1:]
        i+=1
    return filedir

def parseDropDate(data):
    # drop input processing into list of files
    data = data.split(' ')
    data0 = data
    i = 0
    for d in data:
        if not (d.startswith('/') or d.startswith('{')):
            data0[i-1] = data0[i-1]+' '+d
        i+=1
    data =[d.strip('{').strip('}') for d in data0]

    return data

def addInDrop(dropdata):
    global formpaths_to_formdicts
    data =parseDropDate(dropdata)
    app.setStretch('column')
    with app.tabbedFrame('inputs'):
        files = []
        for f in data:
            files = files + getFiles([f],'pdf')
        for f in files:
            f=f.decode('utf-8')
            addIn(f)

def addIn(filepath):
    instoouts[filepath]=[]
    with app.tabbedFrame('inputs'):
        if filepath in app.getAllListItems('inputs'):
            return
        app.addListItem('inputs', filepath)
        # worker_entries = get_forms(f, inputs_to_form_entries)
        with open(filepath,'rb') as infile:
            worker_entries= PdfFileReader(infile).getFields()
            formpaths_to_formdicts[filepath] = worker_entries
            fields = {}
            for k in worker_entries.keys():
                fields[k]= worker_entries[k]['/V']
            formpaths_to_fielddicts[filepath] = fields
            try:
                fullname = fields['FIRST NAME'] + fields['LAST NAME']
            except KeyError:
                fullname = filepath.split('/')[-1]
            except TypeError:
                fullname = filepath.split('/')[-1]

            app.setStretch('both')
            with app.tab(filepath):
                # intab = intabs.widgetStore[filepath][0]
                # print(intab)
                # pprint(dir(intab))
                # intab.bind('<Button-1>', updateInpathFromIntab)
                app.setTabText('inputs',filepath,fullname)
                with app.scrollPane('Worker Form Edit' + filepath):
                    infields =sorted([key.strip(' ') for key in worker_entries.keys()])

                    num_infields = range(len(infields))
                    app.setSticky('e')
                    [app.addLabel(infields[i] +'_' + filepath,infields[i][:25],i,2,0) for i in num_infields]
                    [app.setLabelTooltip(infields[i]+'_'+filepath,infields[i]) for i in num_infields]
                    app.setSticky('w')
                    [app.addEntry(infields[i]+'_'+filepath,i,3,1,0)for i in num_infields]
                    try:
                        [app.setEntry(infields[i]+'_'+filepath,worker_entries[infields[i]]['/V']) for i in num_infields]
                    except TypeError:
                        pass
                    except KeyError:
                        pass


def removeInput():
    selected_form = app.getListBox('inputs')
    app.removeListItem('inputs',selected_form)
    app.deleteTabbedFrameTab('inputs', selected_form[0])
    for w in app.widgetManager.group(app.Widgets.Label):
        if selected_form[0] in w:
            app.widgetManager.get(app.Widgets.Label,w).destroy()
    for w in app.widgetManager.group(app.Widgets.Entry):
        if selected_form[0] in w:
            app.widgetManager.get(app.Widgets.Entry,w).destroy()

def removeOutputTemplate():
    app.deleteTabbedFrameTab('out_templates',app.getTabbedFrameSelectedTab('out_templates'))

def saveFormTemplate():
    global template_form
    global template_fields
    global template_form_path
    writer_copy=PdfFileWriter()
    writer_copy.cloneReaderDocumentRoot(template_form)
    for p in range(writer_copy.getNumPages()):
        page = writer_copy.getPage(p)
        writer_copy.updatePageFormFieldValues(page, template_fields)

    with open(template_form_path.strip('.pdf') + '_copy_TEST.pdf','wb') as TEST:
        writer_copy.write(TEST)

selected_form_template_entry = None

def updateFormTemplateEdit():
    app.setEntry('form_entry_edit',app.getListBox(app.getTabbedFrameSelectedTab('form_templates'))[0])
    # global selected_form_template_entry
    # selected_form_template_entry = app.getListBoxPos('form_templates')[0]


def updateFormTemplates():
    global selected_form_template_entry
    global template_fields

    entry = app.getEntry('form_entry_edit')
    app.selectListItemAtPos('form_templates',selected_form_template_entry)
    old_entry = app.getListBox('form_templates')[0]
    if  old_entry != entry:
        # nested dictionary with title, value info
        formbox = template_fields.pop(old_entry)
        template_fields[entry] = formbox
        app.setListItemAtPos('form_templates',selected_form_template_entry,entry)

def generateOutput():
    sel_inpath = app.getListBox('inputs')[0]

    input_name = sel_inpath.split('/')[-1].replace('.pdf','')
    rulesheet = app.getTabbedFrameSelectedTab('rulesheets')
    rulesheet_name = rulesheet.split('/')[-1].replace('.txt','')
    out_template = app.getTabbedFrameSelectedTab('out_templates')
    out_template_name = out_template.split('/')[-1].replace('.docx','')
    output_name = input_name+'_'+out_template_name+'.docx'
    outpath = output_dirpath+'/'+output_name

    indoc = Document(out_template)
    indoc.save(outpath)
    with app.tabbedFrame('output_preview'):

        instoouts[sel_inpath]+=[outpath]
        outstoins[outpath] = sel_inpath

        # create or open output preview tab
        with app.tab(outpath):
            # outtab = outtabs.getTab(outpath)
            # outtab.bind('<Button-1>', updateOutpathFromOutTab)
            # outtabs.bind('<Button-1>', updateOutpathFromOutTab)

            app.setTabText('output_preview', outpath, output_name)
            # create output text-area, if it does not exist
            try:
                textwidget = app.getTextAreaWidget(outpath)
            except appJar.appjar.ItemLookupError:
                textwidget = app.addScrolledTextArea(outpath)
                app.addListItem('outputs', outpath)

        outdoc = liveDoc(outpath, textwidget)
        outdocs[outpath] = outdoc



class liveDoc(object):
    def __init__(self, outpath, textwidget):
        self.doc = Document(outpath)
        self.outpath = outpath
        self.text = []
        # A list the length of the document's plaintext, where each is a triple
        # of positions (par,run, char)
        # specifying the character's position in the docx object.
        self.txtodocpos = []
        # The tkinter text widget with which the doc is to be binded for live editing.
        self.textwidget = textwidget


        row, i, j, k = 0, 0, 0, 0
        for p in self.doc.paragraphs:
            self.text +=['']
            self.txtodocpos += [[]]
            for r in p.runs:
                self.text[row] += r.text
                for c in r.text:
                    self.txtodocpos[row]  += [(i,j,k)]
                    k += 1
                    # if c == '\n':
                    #     column =0
                    #     row += 1
                    #     self.txtodocpos[row] = []
                k = 0
                j += 1
            j = 0
            i+= 1
        # if p.runs:
            print('par '+str(i-1)+' has text')
            row += 1

        # print(self.txtodocpos)

        def key(event):
            ''' Add one character to document.'''
            print(repr(event.char))
            print(event.type)
            if '@' in event.char:
                print('command')
            if event.char and event.char not in [u'\uf700',u'\uf701',u'\uf702',u'\uf703']:
                # delete selected characters, if any
                try:
                    selection = len(event.widget.get('sel.first', 'sel.last'))
                    print('selection length: '+ str(selection))
                    txpos = self.textwidget.index('sel.first').split('.')
                    backspace(event)
                except TclError:
                    txpos = self.textwidget.index('insert').split('.')
                row, col = int(txpos[0])-1,int(txpos[1])
                print(row, col)
                txtodocpos_row = self.txtodocpos[row]
                print(txtodocpos_row)
                # in case we are adding to the end of a paragraph
                if col == len(txtodocpos_row):
                    # in case the paragraph is empty
                    if col == 0:
                        p, r, c = row, 0,0
                    else:
                        docpos = txtodocpos_row[col-1]
                        p, r, c = docpos[0],docpos[1],docpos[2]+1
                else:
                    docpos = txtodocpos_row[col]
                    p, r, c = docpos[0],docpos[1],docpos[2]
                currun = self.doc.paragraphs[p].runs[r]
                currun.text = currun.text[:c] + event.char +currun.text[c:]
                # insert new pos-list item
                self.txtodocpos[row].insert(col,(p,r,c))
                print(self.txtodocpos[row])
                # adjust the char-pos of all positions left in the run
                # c += 1
                ccol = col
                print(len(currun.text), currun.text)
                while c < len(currun.text):
                    self.txtodocpos[row][ccol] = (p,r,c)
                    print(self.txtodocpos[row])

                    ccol += 1
                    c+=1
                self.text[row] = self.text[row][:col] + event.char + self.text[row][col+1:]

                docpos0 = self.txtodocpos[row]
                print(docpos0)

        # from Scanny: https://github.com/python-openxml/python-docx/issues/33
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

        def backspace(event):
            ''' For cutting, backspacing, deleting '''
            try:
                txpos = self.textwidget.index('sel.last').split('.')
                # delete multi-character selection in text-widget
                selection = len(event.widget.get('sel.first', 'sel.last'))
            except TclError:
                txpos = self.textwidget.index('insert').split('.')
                # or else delete a single character
                selection = 1
            print('selection: ' + str(event.widget.get('sel.first', 'sel.last')))
            print(str(selection)+ ' characters selected')
            row, col = int(txpos[0])-1,int(txpos[1])
            # assuming a new line equates to a new paragraph
            curpar = self.doc.paragraphs[row]
            s = 0
            while s < selection:
                # delete moving backwards from cursor
                while col - 1 not in range(len(self.text[row])):
                    # delete empty paragraph
                    if not self.text[row]:
                        self.text.pop(row)
                        self.txtodocpos.pop(row)
                        delete_paragraph(self.doc.paragraphs[row])
                    row -=1
                    col = len(self.text[row])
                    s+=1


                docpos = self.txtodocpos[row][col-1]
                p, r, c = docpos[0],docpos[1],docpos[2]-1
                currun = self.doc.paragraphs[p].runs[r]
                print(row,col)
                print(self.text[row])
                print(docpos)
                print(currun.text)
                self.text[row]=self.text[row][:col-1] + self.text[row][col:]
                currun.text = currun.text[:c+1]+currun.text[c+2:]
                self.txtodocpos[row]=self.txtodocpos[row][:-1]
                col -= 1
                s+=1
            # print(self.text[row])
            # print(currun.text)

        def paste(event):
            ''' Pastes text on the clipboard into the cursor position.
            '''
            cliptext = event.widget.clipboard_get()
            print('cliptext: '+cliptext)
            if not cliptext and self.textwidget.get('sel.first', 'sel.last'):
                backspace()
            else:
                # write reverse so as to skip moving the cursor
                for c in cliptext[::-1]:
                    event.char = c
                    key(event)
            # try:
            #     txpos = self.textwidget.index('sel.first').split('.')
            #     # deleting selection in text-widget
            #     backspace(event)
            # except TclError:
            #     txpos = self.textwidge.index('insert').split('.')
            # print('welcome to paste')
            # int(txpos[0])-1,int(txpos[1])
            # # TODO: Paste with style/formatting?
            # docpos = self.txtodocpos[txpos]
            # p, r, c = docpos[0],docpos[1],docpos[2]
            # currun = self.doc.paragraphs[p].runs[r]
            # pastetxt = app.clipboard_get()
            # self.text = self.text[:txpos] + pastetxt + self.text[txpos + len(pastetxt):]
            # currun.text = currun.text[c] + pastetxt + currun.text[c + len(pastetxt):]
            # self.txtodocpos = self.txtodocpos[:txpos] + [(p,r,c+i) for i in range(len(pastetxt))] + txtodocpos[txpos + len(pastetxt):]

        textwidget.bind('<KeyRelease>', key)
        textwidget.bind('<BackSpace>', backspace)
        textwidget.bind('<Command-v>', paste)
        textwidget.bind('<Command-c', None)
        textwidget.bind('Command-x',None)

        self.textwidget.insert("insert",'\n'.join(self.text))
        # # Empty event obkject allowing for simulation of keystrokes
        # event = Event()
        # compiled_rules = compileRules()
        # for rule in compiled_rules:
        #     for replacee in rule.replacees:
        #         print('replacee: '+replacee)
        #         print('replacement: '+rule.replacement)
        #         replacement = rule.replacement
        #         # find replacement positions for par in self.text:
        #         for par in self.text:
        #             j = par.upper().find(replacee.upper())
        #             while j >=0:
        #                 # the position of the current replacee in the paragraph
        #                 # print(par)
        #                 print('replacee index: '+str(j))
        #                 print('paragraph length: '+str(len(par)))
        #                 if j>2 and par[j-2] == '.':
        #                     replacement = replacement.capitalize()
        #                 elif j>4 and self.text[j-4:j-1] == '(n)':
        #                     r = replacement[0]
        #                     self.textwidget.mark_set("insert", "%d.%d" % (i + 1, j-1))
        #                     if r == 'a' or r=='e' or r=='u' or r=='i' or r=='o':
        #                         backspace(event)
        #                         self.textwidget.mark_set("insert", "%d.%d" % (i + 1, j-3))
        #                         backspace(event)
        #                         j -= 2
        #                     else:
        #                         backspace(event)
        #                         backspace(event)
        #                         backspace(event)
        #                         j -= 3
        #                 self.textwidget.clipboard_clear()
        #                 self.textwidget.clipboard_append(replacement)
        #                 self.textwidget.mark_set("sel.first", "%d.%d" % (i, j))
        #                 # print('replacee length: '+ str(len(replacee)))
        #                 # # print('sel.first position: '+self.textwidget.get('sel.first'))
        #                 j += len(replacement)
        #                 #
        #                 self.textwidget.mark_set("sel.last", "%d.%d" % (i, j))
        #                 # print('sel.first position: '+self.textwidget.get('sel.first'))
        #                 print('selection: ' + str(self.textwidget.get('sel.first', 'sel.last')))
        #                 event.widget = self.textwidget
        #                 paste(event)
        #                 j = par[j:].upper().find(replacee.upper())
        #
        #             i += 1

    def save(self, outpath):
        return self.doc.save(outpath)

def updateRuleeditEntry():
    rules = app.getListBox(app.getTabbedFrameSelectedTab('rulesheets')+'_rules')
    if rules:
        app.setEntry('rule_edit',rules[0])

def addInTemplatesDrop(dropdata):
    data = parseDropDate(dropdata)
    with app.tabbedFrame('form_templates'):
        files = []
        for f in data:
            files = files + getFiles([f],'pdf')
        for f in files:
            addInTemplate(f)

def addInTemplate(filepath):
    with app.tabbedFrame('form_templates'):
        template_name = filepath.split('/')[-1].replace('.pdf','')
        with open(filepath,'rb') as intemp:
            template_form = PdfFileReader(intemp)
            with app.tab(template_name):
                app.setStretch('both')
                app.addListBox(template_name,sorted(template_form.getFields()),0,0,10,10)
                app.setListBoxGroup(template_name)
                app.setListBoxChangeFunction(template_name,updateFormTemplateEdit)



def addOutTemplatesDrop(dropdata):
    data = parseDropDate(dropdata)
    files = []
    for f in data:
        files = files + getFiles([f],'docx')
    for f in files:
        add_out_template(f)


def addOutTemplate(f):
    global out_temp_paths
    out_temp_paths += [f]
    app.setStretch('both')
    with app.tabbedFrame('out_templates'):
        # Assumes the name (not just path) of each template file is unique.
        template_name = f.split('/')[-1].replace('.docx','')

        with app.tab(f):
            app.setTabText('out_templates',f,template_name)
            # app.addListBox(template_name + '_boxes',boxes,0,2)
            out_temp_listbox = app.addListBox(f ,None,0,1,10,10)
        app.setListBoxGroup(f)
        app.setListBoxChangeFunction(f,updateRuleeditEntry)
        outfields = getDocFields(f)
        outfields = [s.decode().upper().strip(' ') for s in sorted(outfields)]
        for outfield in outfields:
            if outfield not in app.getAllListItems(f):
                app.addListItem(f, outfield)
        app.getListBoxWidget(f).bind("<Double-Button-1>", updateRuleFromOutfield)

        # add empty rules for unaccounted-for outfields
        replacee_lists = []
        sel_rulesheet = app.getTabbedFrameSelectedTab('rulesheets')
        if sel_rulesheet == None:
            path = rulesheet_dirpath+'/'+'New Rules.txt'
            new_rulesheet = open(path, 'wb')
            new_rulesheet.close()
            addRulesheet(path)
        sel_rulesheet = app.getTabbedFrameSelectedTab('rulesheets')
        rules_listbox = sel_rulesheet+'_rules'
        rules = app.getAllListItems(rules_listbox)
        # TODO:
        for rule in rules:
             replacee_lists += box_ptrn.findall(rule.decode().split('replaced by')[0])

        outfields = app.getAllListItems(f)
        for outfield in outfields:
            in_rule = False
            for replacee_list in replacee_lists:
                if outfield in replacee_list:
                    in_rule=True
            if not in_rule:
                rules+= [outfield +' replaced by {}']
        rules = sorted(rules)
        app.clearListBox(rules_listbox)
        app.addListItems(rules_listbox, rules)

def updateRuleFromOutfield(dblclick):
    rules_listbox = app.getTabbedFrameSelectedTab('rulesheets')+'_rules'
    out_temp_path = app.getTabbedFrameSelectedTab('out_templates')
    outfield = app.getListBox(out_temp_path)[0]
    # TODO: limit to rule-outfields using parser
    for rule in app.getAllListItems(rules_listbox):
        if outfield in rule:
            app.selectListItem(rules_listbox, rule)
            break

def getDocFields(docpath):
    # Returns a list of fields to be replaced in a text-based document, where
    # such fields are denoted by brackets "[]" surrounding text describing
    # info to be filled in.
    out_doc = Document(docpath)
    fillable_fields = []
    j=0
    for p in out_doc.paragraphs:
        inline = p.runs
        par_text =''
        # Assembles paragraph-text across differently-styled runs.
        for i in range(len(inline)):
            par_text = par_text + inline[i].text
        boxes = re.findall('(\[.*?\])', par_text)
        for box in boxes:
            fillable_fields = fillable_fields + [box.encode('utf8')]
        j=j+1
    return fillable_fields

def updateRules():
    sel_rulesheet = app.getTabbedFrameSelectedTab('rulesheets')
    sel_rule = app.getListBox(sel_rulesheet+'_rules')
    new_rule = app.getEntry('rule_edit')
    app.setListItemAtPos(sel_rulesheet+'_rules',app.getListBoxPos(sel_rulesheet+'_rules')[0],new_rule)
    app.selectListItem(sel_rulesheet+'_rules', app.getEntry('rule_edit'))

def addCondition():
    rule = app.getEntry('rule_edit')
    if '] when ' in rule:
        subrule = re.search('\] when (\[.*?\] .+? \[.*?\], )+or \[.*?\] .+? \[.*?\]|\] when (\[.*?\] .+? \[.*?\])',rule).group(0)
        new_sub = subrule.replace('], or [','], [') + ', or [] == []'
        new_rule = rule.replace(subrule, new_sub)
    else:
        new_rule = rule + ' when [] == []'
    app.setEntry('rule_edit', new_rule)

def addReplacee():
    rule = app.getEntry(rule_edit)
    p = rule.find(' replaced by')
    app.setEntry(rule_edit, rule[:p]+', {}')

def addReplacement():
    rule = app.getEntry('rule_edit')
    subrule = re.search('(\] replaced by \[.*?\], )+or \[.*?\]|\] replaced by (\[.*?\])',rule).group(0)
    p = rule.find(subrule)
    new_sub = subrule.replace('], or [','], [') + ', or []'
    new_rule = rule.replace(subrule, new_sub)
    app.setEntry('rule_edit', new_rule)
    updateRules()

def pasteEntry2Rule():
    src_entry, trg_entry = 'form_entry_edit','rule_edit'
    # Replaces the first empty box in the rule with the selected form template field.
    app.setEntry(trg_entry,
    re.split('(\{ *\})',app.getEntry(trg_entry),1)[0] + '{'+app.getEntry(src_entry)
                            +'}'+re.split('\{ *\}',app.getEntry(trg_entry),1)[1])
    updateRules()

def delEntryFromRule():
    trg = 'rule_edit'
    app.setEntry(trg, re.sub('\}.+?\{ ', '}{ ',app.getEntry(trg)[::-1],1)[::-1])
    updateRules()

def saveRulesheet():
    rulesheet = app.getTabbedFrameSelectedTab('rulesheets')
    with open(rulesheet, "wb") as text_file:
        txt= ''
        for t in app.getAllListItems(rulesheet+'_rules'):
            txt += t
        text_file.write(txt.strip('\n').encode('utf8'))

def pasteOutfield():
    rule = app.getEntry('rule_edit')
    subrule  = rule.split('] replaced by {')[0] +']'
    outfield = app.getListBox(app.getTabbedFrameSelectedTab('out_templates'))[0]
    if ('], and [' in subrule) or ('] and [' in subrule):
        new_sub = subrule.replace('], and [', '], [').replace('] and [', '], [')
        new_sub = new_sub +', and ' + outfield
    else:
        new_sub = subrule + ' and ' + outfield
    new_rule = rule.replace(subrule, new_sub)
    app.setEntry('rule_edit', new_rule)
    updateRules()

def addRulesheetsDrop(dropdata):
    data = parseDropDate(dropdata)
    files = []
    for f in data:
        files += getFiles(f)
    for f in files:
        addRulesheet(f)

def addRulesheet(path):
    rules = [rule for rule in open(path,'rb').readlines()]
    rules=sorted(rules)
    sheetname = path.split('/')[-1].replace('.txt','')
    with app.tabbedFrame('rulesheets'):
        with app.tab(path):
            app.setTabText('rulesheets',path, sheetname)
            app.addListBox(path+'_rules', [], 0,0,10,10)
            for rule in rules:
                app.addListItem(path+'_rules',rule)
            app.setListBoxGroup(path+'_rules')
            app.setListBoxChangeFunction(path+'_rules',updateRuleeditEntry)

def deleteRule():
    sheetpath = app.getTabbedFrameSelectedTab('rulesheets')
    app.removeListItem(sheetpath+'_rules',app.getListBox(sheetpath+'_rules'))

def updateIntabFromInpath():
    print('updateIntabFromInpath')
    CALLOUT = False
    if app.getListBox('inputs')!= None:
        selected_inpath = app.getListBox('inputs')[0]
        selected_intab = app.getTabbedFrameSelectedTab('inputs')
        if selected_intab != selected_inpath:
            print('CHANGING')
            if CALLOUT:
                app.setTabbedFrameSelectedTab('inputs', selected_inpath)
            else:
                app.setTabbedFrameSelectedTab('inputs', selected_inpath)
    selected_outpaths = app.getListBox('outputs')
    selected_intab = app.getTabbedFrameSelectedTab('inputs')
    CALLOUT = False
    print('desired outpaths:' + str(instoouts[selected_intab]))
    for outpath in selected_outpaths:
        print('actual outpath: '+outpath)
        if outpath not in instoouts[selected_intab]:
            CALLOUT = True
    if CALLOUT:
        print('updating outpath from inpath!')
        updateOutpathFromInpath()

# Catches loop
def updateInpathFromIntab(event):
    # for line in traceback.format_stack():
    #     print(line.strip())
    print('updateInpathFromIntab')
    selected_inpath = app.getListBox('inputs')[0]
    selected_intab = app.getTabbedFrameSelectedTab('inputs')

    if not selected_intab == selected_inpath:
        # print('CHANGING')
        app.selectListItem('inputs',selected_intab)


def updateOutTabFromOutpath():
    print('updateOutTabFromOutpath')
    outpaths = app.getListBox('outputs')
    outpath = outpaths[0]
    inpath = outstoins[outpath]
    if app.getTabbedFrameSelectedTab('output_preview') not in outpaths:
        # print('CHANGING')
        app.setTabbedFrameSelectedTab('output_preview', outpaths[0])

    inpath = outstoins[outpath]
    if inpath != app.getListBox('inputs')[0]:
        print('desired inpath: ' + inpath)
        print('actual inpath: ' + app.getListBox('inputs')[0])
        print('! updating inpath from outpath !')
        updateInpathFromOutpath()

# catches loop
def updateOutpathFromOutTab(event):
    print('updateOutpathFromOutTab')
    outpath = app.getTabbedFrameSelectedTab('output_preview')
    if outpath not in app.getListBox('outputs'):
        # print('CHANGING')
        lb = app.getListBoxWidget('outputs')
        END = len(app.getAllListItems('outputs'))
        lb.selection_clear(0, END)
        app.selectListItem('outputs',outpath)


def updateOutpathFromInpath():
    print('updateOutpathFromInpath')
    inpath = app.getListBox('inputs')[0]
    lb = app.getListBoxWidget('outputs')
    END = len(app.getAllListItems('outputs'))
    lb.selection_clear(0, END)
    for outpath in instoouts[inpath]:
        app.selectListItem('outputs',outpath)

def updateInpathFromOutpath():
    print('updateInpathFromOutpath')
    outpath = app.getListBox('outputs')[0]
    inpath = outstoins[outpath]
    if app.getListBox('inputs')[0] != inpath:
        # print('CHANGING')
        app.selectListItem('inputs', inpath)



with gui("OPC form2doc") as app:
    with app.panedFrameVertical('inputs'):
        with app.panedFrame('input_paths'):
            app.setStretch('column')
            app.addLabel('Inputs', 'Inputs',0,0,1,1)
            app.addButton('Remove Input',removeInput,32,0)
            app.setButtonTooltip('Remove Input', 'Does not delete input form from hard drive.')

            app.setStretch('both')
            app.addListBox('inputs',[],1,0,1,31)
            app.setListBoxGroup('inputs')
            app.setListBoxChangeFunction('inputs', updateIntabFromInpath)

            with app.panedFrame('input_tabs'):
                app.setStretch('column')
                app.addLabel('Input Preview', 'Input Preview', 0,1)
                app.addButton('Save Form',saveFormedit,32,1)
                app.setStretch('both')
                with app.tabbedFrame('inputs',1,1,1,31) as intabs:
                    app.setListBoxDropTarget('inputs', addInDrop, replace=False)
                    app.addLabel('plc','plc')
                    app.hideLabel('plc')
                    app.setTabbedFrameChangeFunction('inputs',updateInpathFromIntab)
                    intabs.changeOnFocus = False
            with app.panedFrame('outpaths'):
                app.setStretch('column')
                app.addLabel('Outputs','Outputs',0,2)
                app.addButton('Generate Output',generateOutput,32,2)
                app.setStretch('both')
                app.addListBox('outputs',[],1,2,1,31)
                app.setListBoxGroup('outputs')
                app.setListBoxMulti('outputs')
                app.setListBoxChangeFunction('outputs', updateOutTabFromOutpath)
            with app.panedFrame('output_previews'):
                app.setStretch('column')
                app.addLabel('Output Preview','Output Preview',0,3)
                app.addButton('Save Output',saveOutput,32,3)
                app.setStretch('both')
                with app.tabbedFrame('output_preview',1,3) as outtabs:
                    app.addLabel('plc1','plc')
                    app.hideLabel('plc1')

                    app.setTabbedFrameChangeFunction('output_preview', updateOutpathFromOutTab)
                    outtabs.changeOnFocus = False
                    # outtabs.bind('<Button-1>', updateOutpathFromOutTab)

        with app.panedFrame('templates'):
            with app.panedFrame('in_temps'):
                app.setStretch('both')
                with app.tabbedFrame('form_templates',35,0,1,30):
                    app.setTabbedFrameDropTarget('form_templates', addInTemplatesDrop)
                app.setStretch('column')
                app.addLabel('Form Template','Input Templates',34,0)

                app.addButton('>',pasteEntry2Rule,67,0)
                app.addButton('Save to Form Template', saveFormTemplate,68,0)
                app.addEntry('form_entry_edit',66,0)

            with app.panedFrame('rulesheets'):
                app.setStretch('column')
                app.addLabel('Rule Sheets','Rule Sheets',34,1,2,1)

                app.addButton('Add Replacement',addReplacement,68,1)
                app.addButton('Save Rule Sheet',saveRulesheet,69,2)
                app.addButton('Delete Rule', deleteRule,69,1)
                app.addEntry('rule_edit',66,1,2,1)
                app.setEntrySubmitFunction('rule_edit',updateRules)
                app.addButton('delete last',delEntryFromRule,67,1)
                app.addButton('Add Condition',addCondition,68,2)

                app.setStretch('both')
                with app.tabbedFrame('rulesheets',35,1,2,30):
                    app.setTabbedFrameDropTarget('rulesheets', addRulesheetsDrop)



            with app.panedFrame('out_temps'):
                app.setStretch('column')
                app.addLabel('out_templates','Output Templates',34,3)
                app.addButton('<', pasteOutfield,66,3)

                app.setStretch('both')
                with app.tabbedFrame('out_templates',35,3,10,10):
                    app.setTabbedFrameDropTarget('out_templates', addOutTemplatesDrop)

                app.addButton('Remove Output Template',removeOutputTemplate,69,3)
                app.setButtonTooltip('Remove Output Template','Does not delete output template from hard drive.')




    # Loading default templates and rules
    for path in os.listdir(rulesheet_dirpath):
        if path.endswith('.txt'):
            addRulesheet(rulesheet_dirpath+'/'+path)
    for path in os.listdir(input_templates_dirpath):
        if path.endswith('.pdf'):
            addInTemplate(input_templates_dirpath+'/'+path)
    for path in os.listdir(output_templates_dirpath):
        if path.endswith('.docx'):
            addOutTemplate(output_templates_dirpath+'/'+path)
    # Load default input folder
    for path in os.listdir(input_dirpath):
        if path.endswith('.pdf'):
            addIn(input_dirpath+'/'+path)

    # app.setEntrySubmitFunction('form_entry_edit',updateFormTemplates)

    for inpath in app.getAllListItems('inputs'):
        app.selectListItem('inputs',inpath)
        for templatepath in out_temp_paths:
            app.setTabbedFrameSelectedTab('out_templates', templatepath)
            generateOutput()


    # app.setFont(15)
    # app.setBg("black")
    # app.setFg("lightGray")



app.go()
